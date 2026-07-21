import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Inicializar documento
doc = docx.Document()

# Título Principal
title = doc.add_heading('📄 RESUMEN EJECUTIVO: GUÍA DE DESARROLLO Y DESPLIEGUE', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Dashboard Ejecutivo de Ventas con Python, Streamlit y GitHub\n')

# Sección 1
doc.add_heading('1. Programas que debes instalar', level=1)
p1 = doc.add_paragraph()
p1.add_run('• Python (v3.10+): ').bold = True
p1.add_run('Lenguaje base. Descargar de https://www.python.org/downloads/ (Marcar "Add Python to PATH").\n')
p1.add_run('• Visual Studio Code: ').bold = True
p1.add_run('Editor de código. Descargar de https://code.visualstudio.com/\n')
p1.add_run('• Git & Git Bash: ').bold = True
p1.add_run('Consola de control de versiones. Descargar de https://git-scm.com/downloads')

# Sección 2
doc.add_heading('2. Librerías e Instalación', level=1)
doc.add_paragraph('Para instalar todas las librerías necesarias, ejecuta en Git Bash o la Terminal:')
doc.add_paragraph('pip install streamlit pandas plotly openpyxl python-docx', style='Quote')

# Sección 3
doc.add_heading('3. Registro de Cuentas', level=1)
doc.add_paragraph('• GitHub: Crear cuenta en https://github.com/signup\n• Streamlit Cloud: Ingresar en https://share.streamlit.io e iniciar sesión haciendo clic en "Continue with GitHub".')

# Sección 4
doc.add_heading('4. Sincronización Local y Despliegue', level=1)
doc.add_paragraph('1. Probar localmente: streamlit run app.py\n2. Subir cambios a GitHub:\n   git add .\n   git commit -m "Actualizacion"\n   git push origin main\n3. Desplegar en la nube en share.streamlit.io vinculando el repositorio.')

# Sección 5
doc.add_heading('5. Prompt Maestro para el Dashboard', level=1)
doc.add_paragraph('Utiliza el siguiente prompt en la IA para reconstruir la app:')
doc.add_paragraph(
    "Actúa como un Desarrollador Senior de Data Analytics. Escribe un código en Streamlit y Plotly para un 'Executive Sales Dashboard'. "
    "Debe incluir lectura de Excel multipestaña, corrección de fechas latinoamericanas (dayfirst=True), filtros jerárquicos (fechas/meses/categorías), "
    "Métricas de Ventas, Transacciones y Ticket Promedio, Gráfico de Líneas Spline, Dona de Categorías y Lollipop Top 10 con temática Dark Executive.",
    style='Quote'
)

# Guardar documento
nombre_archivo = "Resumen_Ejecutivo_Dashboard.docx"
doc.save(nombre_archivo)
print(f"✅ ¡El archivo '{nombre_archivo}' se ha generado exitosamente en tu carpeta!")